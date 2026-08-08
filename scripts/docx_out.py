#!/usr/bin/env python3
"""DOCX emitter: the editable deliverable, mirroring the solved PDF geometry.

The fit is solved once, in `render.py`, against Carlito metrics. This module
reproduces that solution in Word by using the same margins, sizes, spacings and
the scale the solver picked, and by naming Calibri as the document font, which
is metric-compatible with Carlito. Same metrics plus same geometry means Word
breaks lines where the measured layout broke them, so the page count survives
the trip.

Word still owns its own renderer, so treat the PDF as the proof of fit and the
DOCX as the editable copy of the same document.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Calibri for Office (metric-identical to the Carlito the PDF embeds); Carlito
# is the fallback so Linux and LibreOffice pick the same metrics.
FONT = "Calibri"
FONT_FALLBACK = "Carlito"

ACCENT = "2B4C7E"
DARK = "1A1A1A"
GRAY = "555555"


def _set_font(run, size, bold=False, color=DARK, charspace_pt=0.0):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)
    if charspace_pt:
        spacing = OxmlElement("w:spacing")  # character spacing, in twentieths of a pt
        spacing.set(qn("w:val"), str(int(round(charspace_pt * 20))))
        rPr.append(spacing)


def _spacing(paragraph, before=0.0, after=0.0, line_pt=None):
    """Spacing in points. `line_pt` is EXACT leading, not a multiple.

    Word reads a line-spacing multiple as a multiple of the font's own line
    height (~1.22em for Calibri), not of the point size, which inflates every
    line ~22% past the measured layout and pushes a full page onto a second
    one. An exact value in points is the only way the DOCX keeps the page count
    the PDF was solved for.
    """
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_pt:
        pf.line_spacing = Pt(line_pt)


def _setup_page(doc):
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.5), Inches(11)
    s.top_margin = s.bottom_margin = Inches(0.5)
    s.left_margin = s.right_margin = Inches(0.7)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rFonts.set(qn(attr), FONT)
    rFonts.set(qn("w:eastAsia"), FONT_FALLBACK)


def _setup_bullets(doc) -> str:
    """Define a bullet list and return its numId."""
    numbering = doc.part.numbering_part._element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "1")
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    for tag, val in (("w:start", "1"), ("w:numFmt", "bullet"),
                     ("w:lvlText", "•"), ("w:lvlJc", "left")):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        lvl.append(el)
    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "216")     # 0.15" indent, matching the PDF's bullet inset
    ind.set(qn("w:hanging"), "216")
    pPr.append(ind)
    lvl.append(pPr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "1")
    ref = OxmlElement("w:abstractNumId")
    ref.set(qn("w:val"), "1")
    num.append(ref)
    numbering.append(num)
    return "1"


def _bullet(doc, text, ctx, num_id):
    p = doc.add_paragraph()
    _set_font(p.add_run(text), ctx["body"], color=DARK)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    for tag, val in (("w:ilvl", "0"), ("w:numId", num_id)):
        el = OxmlElement(tag)
        el.set(qn("w:val"), val)
        numPr.append(el)
    pPr.append(numPr)
    _spacing(p, after=ctx["sp_after_bullet"], line_pt=ctx["body"] * ctx["lead"])


def _section(doc, title, ctx):
    p = doc.add_paragraph()
    _set_font(p.add_run(title.upper()), ctx["section"], bold=True, color=ACCENT)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    for attr, val in (("w:val", "single"), ("w:sz", "4"), ("w:space", "2"), ("w:color", ACCENT)):
        bottom.set(qn(attr), val)
    pBdr.append(bottom)
    pPr.append(pBdr)
    _spacing(p, before=ctx["sp_before_section"], after=ctx["sp_after_section"],
             line_pt=ctx["section"] * ctx["lead"])


def _entry_head(doc, title, trailer, ctx, trailer_color=GRAY):
    p = doc.add_paragraph()
    _set_font(p.add_run(title), ctx["body"] + 0.5 * ctx["scale"], bold=True, color=DARK)
    if trailer:
        _set_font(p.add_run(f"  |  {trailer}"), ctx["body"], color=trailer_color)
    _spacing(p, before=ctx["sp_before_entry"], after=0,
             line_pt=(ctx["body"] + 0.5 * ctx["scale"]) * ctx["lead"])
    return p


def _meta_line(doc, bits, ctx):
    bits = [b for b in bits if b]
    if not bits:
        return
    p = doc.add_paragraph()
    _set_font(p.add_run("  |  ".join(bits)), ctx["meta_small"], color=GRAY)
    _spacing(p, after=ctx["sp_after_entry_meta"], line_pt=ctx["meta_small"] * ctx["lead"])


def _daterange(entry) -> str:
    start, end = entry.get("start"), entry.get("end")
    if start and end:
        return f"{start} – {end}"
    return start or end or ""


def render_docx(data: dict, ctx: dict, out_path):
    doc = Document()
    _setup_page(doc)
    doc.add_paragraph("", style="List Bullet")
    doc._body._body.remove(doc.paragraphs[-1]._p)
    num_id = _setup_bullets(doc)

    meta = data.get("meta", {})

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p.add_run(meta.get("name", "").upper()), ctx["name"], bold=True,
              color=ACCENT, charspace_pt=1.4 * ctx["scale"])
    _spacing(p, after=ctx["sp_after_name"], line_pt=ctx["name"] * ctx["lead"])

    if meta.get("title"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run(meta["title"]), ctx["headline"], color=GRAY)
        _spacing(p, after=ctx["sp_after_headline"], line_pt=ctx["headline"] * ctx["lead"])

    parts = [meta[k] for k in
             ("location", "phone", "email", "linkedin", "github", "website", "relocation")
             if meta.get(k)]
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_font(p.add_run("  ·  ".join(parts)), ctx["contact"], color=GRAY)
        _spacing(p, after=ctx["sp_after_contact"], line_pt=ctx["contact"] * ctx["lead"])

    if data.get("summary"):
        p = doc.add_paragraph()
        _set_font(p.add_run(data["summary"]), ctx["body"], color=DARK)
        _spacing(p, after=ctx["sp_after_summary"], line_pt=ctx["body"] * ctx["lead"])

    if data.get("experience"):
        _section(doc, "Experience", ctx)
        for job in data["experience"]:
            _entry_head(doc, job.get("title", ""), job.get("company"), ctx)
            _meta_line(doc, [job.get("location"), _daterange(job)], ctx)
            for b in job.get("bullets", []):
                _bullet(doc, b, ctx, num_id)

    if data.get("skills"):
        _section(doc, "Skills", ctx)
        for label, items in data["skills"].items():
            if not items:
                continue
            p = doc.add_paragraph()
            _set_font(p.add_run(f"{label.replace('_', ' ').title()}: "), ctx["body"],
                      bold=True, color=ACCENT)
            body = ", ".join(items) if isinstance(items, list) else str(items)
            _set_font(p.add_run(body), ctx["body"], color=DARK)
            _spacing(p, after=ctx["sp_after_skill"], line_pt=ctx["body"] * ctx["lead"])

    if data.get("projects"):
        _section(doc, "Projects", ctx)
        for proj in data["projects"]:
            tech = proj.get("tech")
            tech_s = ", ".join(tech) if isinstance(tech, list) else (tech or "")
            _entry_head(doc, proj.get("name", ""), tech_s, ctx)
            _spacing(doc.paragraphs[-1], before=ctx["sp_before_entry"],
                     after=ctx["sp_after_entry_meta"],
                     line_pt=(ctx["body"] + 0.5 * ctx["scale"]) * ctx["lead"])
            if proj.get("description"):
                _bullet(doc, proj["description"], ctx, num_id)

    if data.get("education"):
        _section(doc, "Education", ctx)
        for edu in data["education"]:
            p = doc.add_paragraph()
            _set_font(p.add_run(edu.get("degree", "")), ctx["body"] + 0.5 * ctx["scale"],
                      bold=True, color=DARK)
            if edu.get("honors"):
                _set_font(p.add_run(f"  |  {edu['honors']}"), ctx["body"], color=ACCENT)
            _spacing(p, before=ctx["sp_before_entry"], after=0,
                     line_pt=(ctx["body"] + 0.5 * ctx["scale"]) * ctx["lead"])
            _meta_line(doc, [edu.get("institution"), edu.get("location"), _daterange(edu)], ctx)

    certs = data.get("certifications") or []
    if certs:
        _section(doc, "Certifications", ctx)
        flat = []
        for cert in certs:
            if isinstance(cert, str):
                flat.append(cert)
            else:
                extra = ", ".join(str(x) for x in (cert.get("issuer"),
                                                   cert.get("year") or cert.get("date")) if x)
                flat.append(f"{cert.get('name', '')} ({extra})" if extra else cert.get("name", ""))
        p = doc.add_paragraph()
        _set_font(p.add_run("  ·  ".join(flat)), ctx["body"], color=DARK)
        _spacing(p, after=0, line_pt=ctx["body"] * ctx["lead"])

    doc.save(str(out_path))
    return out_path
